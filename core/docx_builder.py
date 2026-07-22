"""
DOCX构建器 - 将解析的内容构建为DOCX文件
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from typing import List, Optional
from core.parser import PageContent, TextBlock, ImageBlock, TableBlock
import logging

logger = logging.getLogger(__name__)


class DocxBuilder:
    """DOCX文档构建器"""
    
    def __init__(self, output_path: str):
        """
        初始化构建器
        Args:
            output_path: 输出文件路径
        """
        self.output_path = output_path
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Set document styles"""
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        rPr = style._element.get_or_add_rPr()
        rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

        # Set heading styles
        for i in range(1, 4):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(18 - i * 2)
            heading_style.font.bold = True
            rPr = heading_style._element.get_or_add_rPr()
            rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    
    def _is_heading(self, text_block: TextBlock) -> int:
        """
        判断文字块是否为标题，返回标题级别（0表示不是标题）
        """
        # 根据字体大小和是否粗体判断
        if text_block.is_bold and text_block.font_size >= 14:
            if text_block.font_size >= 18:
                return 1
            elif text_block.font_size >= 16:
                return 2
            else:
                return 3
        
        # 根据文字内容特征判断
        text = text_block.text.strip()
        if len(text) < 50 and not text.endswith(('。', '，', '；', '、')):
            # 短句且不以标点结尾，可能是标题
            if text_block.font_size >= 14:
                return 2
        
        return 0
    
    def _add_text_block(self, text_block: TextBlock):
        """添加文字块"""
        heading_level = self._is_heading(text_block)
        
        if heading_level > 0:
            # 添加为标题
            p = self.doc.add_heading(text_block.text, level=heading_level)
        else:
            # 添加为普通段落
            p = self.doc.add_paragraph(text_block.text)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
    
    def _add_image_block(self, image_block: ImageBlock, 
                         ocr_text: str = "",
                         ai_description: str = "",
                         max_width: float = 15.0):
        """
        添加图片块
        Args:
            image_block: 图片块
            ocr_text: OCR识别出的文字
            ai_description: AI生成的描述
            max_width: 最大宽度（厘米）
        """
        if not os.path.exists(image_block.image_path):
            logger.warning(f"图片不存在: {image_block.image_path}")
            return
        
        try:
            # 添加图片
            from PIL import Image
            with Image.open(image_block.image_path) as img:
                width, height = img.size
                
                # 计算缩放比例
                max_width_cm = Cm(max_width)
                aspect_ratio = height / width
                
                # 转换为厘米（假设96 DPI）
                width_cm = width / 96 * 2.54
                height_cm = height / 96 * 2.54
                
                if width_cm > max_width:
                    width_cm = max_width
                    height_cm = width_cm * aspect_ratio
                
                # 插入图片
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_block.image_path, width=Cm(width_cm))
            
            # 添加AI描述（增强模式）
            if ai_description:
                desc_para = self.doc.add_paragraph()
                desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = desc_para.add_run(ai_description)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 139)
                desc_para.paragraph_format.left_indent = Cm(0.5)
                desc_para.paragraph_format.right_indent = Cm(0.5)
            
            # 添加OCR识别出的文字（直接嵌入，无标签前缀）
            if ocr_text:
                for line in ocr_text.split('\n'):
                    if line.strip():
                        ocr_para = self.doc.add_paragraph()
                        ocr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = ocr_para.add_run(line)
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(80, 80, 80)
                        ocr_para.paragraph_format.space_after = Pt(2)
            
            # 添加空行分隔
            self.doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"添加图片失败: {e}")
    
    def _add_table_block(self, table_block: TableBlock):
        """
        添加表格块
        Args:
            table_block: 表格块
        """
        if not table_block.data or len(table_block.data) == 0:
            return
        
        try:
            # 检查表格数据有效性
            max_cols = max(len(row) for row in table_block.data) if table_block.data else 0
            if max_cols == 0:
                return
            
            # 创建表格
            table = self.doc.add_table(rows=len(table_block.data), cols=max_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 填充数据
            for i, row_data in enumerate(table_block.data):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    if j < max_cols:
                        cell = row.cells[j]
                        cell.text = str(cell_text) if cell_text is not None else ""
                        
                        # 设置单元格字体
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            
            # 设置表头样式（第一行）
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
            
            # 添加空行
            self.doc.add_paragraph()
            
        except Exception as e:
            logger.error(f"添加表格失败: {e}")
            # 如果表格创建失败，尝试添加为图片（如果是图片表格）
            if table_block.is_image and os.path.exists(table_block.image_path):
                img_block = ImageBlock(
                    image_path=table_block.image_path,
                    x=0, y=0, width=400, height=300
                )
                self._add_image_block(img_block, ocr_text="表格识别失败，保留原图")
    
    def build(self, pages: List[PageContent], 
              ocr_results: dict = None,
              descriptions: dict = None):
        """
        构建DOCX文档
        Args:
            pages: 页面内容列表
            ocr_results: OCR结果 {image_path: text}
            descriptions: AI描述结果 {image_path: description}
        """
        if ocr_results is None:
            ocr_results = {}
        if descriptions is None:
            descriptions = {}
        
        for page in pages:
            # 合并所有内容块并按Y坐标排序
            all_blocks = []
            
            for tb in page.text_blocks:
                all_blocks.append(('text', tb.y, tb))
            
            for ib in page.image_blocks:
                all_blocks.append(('image', ib.y, ib))
            
            for tbl in page.table_blocks:
                all_blocks.append(('table', tbl.y, tbl))
            
            # 按Y坐标排序（从上到下）
            all_blocks.sort(key=lambda x: x[1])
            
            # 按顺序添加内容
            for block_type, _, block in all_blocks:
                if block_type == 'text':
                    self._add_text_block(block)
                elif block_type == 'image':
                    ocr_text = ocr_results.get(block.image_path, "")
                    desc = descriptions.get(block.image_path, "")
                    self._add_image_block(block, ocr_text, desc)
                elif block_type == 'table':
                    self._add_table_block(block)
        
        # 保存文档
        self.doc.save(self.output_path)
        logger.info(f"DOCX文档已保存: {self.output_path}")
    
    def add_metadata(self, title: str = "", author: str = "", 
                     description: str = ""):
        """添加文档元数据"""
        if title:
            self.doc.core_properties.title = title
        if author:
            self.doc.core_properties.author = author
        if description:
            self.doc.core_properties.comments = description
