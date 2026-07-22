"""
文件解析器 - 支持PDF、DOCX、XLSX
"""
import os
import io
import fitz  # PyMuPDF
from PIL import Image
from typing import List, Dict, Tuple, Optional, Any
from dataclasses import dataclass
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


@dataclass
class TextBlock:
    """文字块"""
    text: str
    x: float
    y: float
    width: float
    height: float
    font_size: float = 12
    is_bold: bool = False
    page_num: int = 0


@dataclass
class ImageBlock:
    """图片块"""
    image_path: str
    x: float
    y: float
    width: float
    height: float
    page_num: int = 0
    caption: str = ""  # 图注


@dataclass
class TableBlock:
    """表格块"""
    data: List[List[str]]  # 二维表格数据
    x: float
    y: float
    width: float
    height: float
    page_num: int = 0
    is_image: bool = False  # 是否是从图片识别的表格
    image_path: str = ""  # 如果是图片表格，保存图片路径


@dataclass
class PageContent:
    """页面内容"""
    page_num: int
    text_blocks: List[TextBlock]
    image_blocks: List[ImageBlock]
    table_blocks: List[TableBlock]
    width: float
    height: float


class DocumentParser:
    """文档解析器"""
    
    def __init__(self, temp_dir: str):
        self.temp_dir = temp_dir
        os.makedirs(temp_dir, exist_ok=True)
    
    def parse(self, file_path: str) -> List[PageContent]:
        """
        解析文档
        Args:
            file_path: 文件路径
        Returns:
            页面内容列表
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return self._parse_docx(file_path)
        elif ext in ['.xls', '.xlsx']:
            return self._parse_xlsx(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _parse_pdf(self, file_path: str) -> List[PageContent]:
        """解析PDF文件"""
        pages = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_idx in range(len(doc)):
                page = doc[page_idx]
                page_width = page.rect.width
                page_height = page.rect.height
                
                text_blocks = []
                image_blocks = []
                table_blocks = []
                
                # 提取文字块
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if block["type"] == 0:  # 文字块
                        for line in block["lines"]:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            
                            if line_text.strip():
                                # Use max font size in line and check bold flags
                                spans = line["spans"]
                                max_size = max((s["size"] for s in spans), default=12)
                                is_bold = any(
                                    "bold" in s["font"].lower() or (s.get("flags", 0) & 16)
                                    for s in spans
                                ) if spans else False
                                text_blocks.append(TextBlock(
                                    text=line_text,
                                    x=block["bbox"][0],
                                    y=block["bbox"][1],
                                    width=block["bbox"][2] - block["bbox"][0],
                                    height=block["bbox"][3] - block["bbox"][1],
                                    font_size=max_size,
                                    is_bold=is_bold,
                                    page_num=page_idx
                                ))
                
                # 提取图片
                image_list = page.get_images(full=True)
                for img_idx, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # 保存图片
                        img_filename = f"page_{page_idx}_img_{img_idx}.{image_ext}"
                        img_path = os.path.join(self.temp_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(image_bytes)
                        
                        # 获取图片在页面中的位置
                        rect = None
                        try:
                            rects = page.get_image_rects(xref)
                            if rects:
                                rect = rects[0]
                        except Exception:
                            pass
                        
                        if rect:
                            image_blocks.append(ImageBlock(
                                image_path=img_path,
                                x=rect.x0,
                                y=rect.y0,
                                width=rect.width,
                                height=rect.height,
                                page_num=page_idx
                            ))
                        else:
                            # 如果找不到位置，放在页面底部
                            image_blocks.append(ImageBlock(
                                image_path=img_path,
                                x=0,
                                y=page_height - 200,
                                width=200,
                                height=200,
                                page_num=page_idx
                            ))
                    except Exception as e:
                        logger.warning(f"提取PDF图片失败: {e}")
                
                # 尝试识别表格（通过检测线条或特殊布局）
                # 这里使用简单的启发式方法：检测页面中的表格区域
                tables = page.find_tables()
                if tables:
                    for table_idx, table in enumerate(tables):
                        try:
                            table_data = table.extract()
                            if table_data:
                                rect = table.bbox
                                table_blocks.append(TableBlock(
                                    data=table_data,
                                    x=rect.x0,
                                    y=rect.y0,
                                    width=rect.width,
                                    height=rect.height,
                                    page_num=page_idx,
                                    is_image=False
                                ))
                        except Exception as e:
                            logger.warning(f"提取PDF表格失败: {e}")
                
                pages.append(PageContent(
                    page_num=page_idx,
                    text_blocks=text_blocks,
                    image_blocks=image_blocks,
                    table_blocks=table_blocks,
                    width=page_width,
                    height=page_height
                ))
            
            doc.close()
            
        except Exception as e:
            logger.error(f"解析PDF失败: {e}")
            raise
        
        return pages
    
    def _parse_docx(self, file_path: str) -> List[PageContent]:
        """解析DOCX文件"""
        from docx import Document
        
        pages = []
        text_blocks = []
        image_blocks = []
        table_blocks = []
        
        try:
            doc = Document(file_path)
            
            # 处理老格式DOC文件 - 转换为DOCX
            if file_path.endswith('.doc'):
                logger.info("检测到DOC格式，尝试转换...")
                # 这里需要依赖外部工具如LibreOffice，先简单处理
                # 实际使用时建议在安装脚本中安装LibreOffice
                pass
            
            # 提取段落文字
            y_position = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    text_blocks.append(TextBlock(
                        text=para.text,
                        x=0,
                        y=y_position,
                        width=500,
                        height=20,
                        font_size=12,
                        is_bold=para.runs[0].bold if para.runs else False,
                        page_num=0
                    ))
                    y_position += 25
            
            # 提取图片
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    try:
                        image = rel.target_part
                        image_bytes = image.blob
                        image_ext = image.content_type.split('/')[-1]
                        if image_ext == 'jpeg':
                            image_ext = 'jpg'
                        
                        img_filename = f"docx_img_{len(image_blocks)}.{image_ext}"
                        img_path = os.path.join(self.temp_dir, img_filename)
                        with open(img_path, "wb") as f:
                            f.write(image_bytes)
                        
                        image_blocks.append(ImageBlock(
                            image_path=img_path,
                            x=0,
                            y=y_position,
                            width=400,
                            height=300,
                            page_num=0
                        ))
                        y_position += 310
                    except Exception as e:
                        logger.warning(f"提取DOCX图片失败: {e}")
            
            # 提取表格
            for table in doc.tables:
                try:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            row_data.append(cell.text)
                        table_data.append(row_data)
                    
                    if table_data:
                        table_blocks.append(TableBlock(
                            data=table_data,
                            x=0,
                            y=y_position,
                            width=500,
                            height=len(table_data) * 25,
                            page_num=0,
                            is_image=False
                        ))
                        y_position += len(table_data) * 25 + 10
                except Exception as e:
                    logger.warning(f"提取DOCX表格失败: {e}")
            
            pages.append(PageContent(
                page_num=0,
                text_blocks=text_blocks,
                image_blocks=image_blocks,
                table_blocks=table_blocks,
                width=612,
                height=max(y_position, 792)
            ))
            
        except Exception as e:
            logger.error(f"解析DOCX失败: {e}")
            raise
        
        return pages
    
    def _parse_xlsx(self, file_path: str) -> List[PageContent]:
        """解析XLSX文件"""
        from openpyxl import load_workbook
        
        pages = []
        text_blocks = []
        image_blocks = []
        table_blocks = []
        
        try:
            wb = load_workbook(file_path, data_only=True)
            
            y_position = 0
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # 添加工作表标题
                text_blocks.append(TextBlock(
                    text=f"工作表: {sheet_name}",
                    x=0,
                    y=y_position,
                    width=500,
                    height=25,
                    font_size=14,
                    is_bold=True,
                    page_num=0
                ))
                y_position += 30
                
                # 提取表格数据
                table_data = []
                for row in sheet.iter_rows(values_only=True):
                    row_data = []
                    for cell in row:
                        if cell is not None:
                            row_data.append(str(cell))
                        else:
                            row_data.append("")
                    table_data.append(row_data)
                
                if table_data:
                    table_blocks.append(TableBlock(
                        data=table_data,
                        x=0,
                        y=y_position,
                        width=500,
                        height=len(table_data) * 20,
                        page_num=0,
                        is_image=False
                    ))
                    y_position += len(table_data) * 20 + 20
            
            # 尝试提取图片
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                if hasattr(sheet, '_images') and sheet._images:
                    for img_idx, img in enumerate(sheet._images):
                        try:
                            image_bytes = img._data()
                            img_filename = f"xlsx_img_{sheet_name}_{img_idx}.png"
                            img_path = os.path.join(self.temp_dir, img_filename)
                            with open(img_path, "wb") as f:
                                f.write(image_bytes)
                            
                            image_blocks.append(ImageBlock(
                                image_path=img_path,
                                x=0,
                                y=y_position,
                                width=400,
                                height=300,
                                page_num=0
                            ))
                            y_position += 310
                        except Exception as e:
                            logger.warning(f"提取XLSX图片失败: {e}")
            
            pages.append(PageContent(
                page_num=0,
                text_blocks=text_blocks,
                image_blocks=image_blocks,
                table_blocks=table_blocks,
                width=612,
                height=max(y_position, 792)
            ))
            
        except Exception as e:
            logger.error(f"解析XLSX失败: {e}")
            raise
        
        return pages
